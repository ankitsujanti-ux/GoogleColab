"""Generate Project_Development_Details.docx from codebase knowledge."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def main():
    doc = Document()
    doc.add_heading('Medication Adherence Dashboard – Project Development Details', 0)

    # 1. Project overview
    add_heading(doc, '1. Project Overview', level=1)
    add_para(doc, 'The Medication Adherence Dashboard is an AI-driven application for monitoring patient medication adherence. It provides real-time analytics, risk assessment, automated interventions (notifications via SMS, email, voice call, WhatsApp, Pushover), and an AI chatbot for patient support.')
    add_para(doc, 'Key features:')
    add_bullet(doc, 'Real-time dashboard with KPIs and charts (Plotly)')
    add_bullet(doc, 'AI-powered risk assessment and patient monitoring using OpenAI GPT')
    add_bullet(doc, 'Automated patient notifications across multiple channels')
    add_bullet(doc, 'WebSocket-based real-time updates (Flask-SocketIO)')
    add_bullet(doc, 'Agentic AI pipeline: data quality, risk, routing, throttling, orchestration, sentiment, care routing, and chatbot agents')

    # 2. Frameworks and technologies
    add_heading(doc, '2. Frameworks and Technologies Used', level=1)

    add_heading(doc, '2.1 Backend', level=2)
    add_bullet(doc, 'Python 3.8+')
    add_bullet(doc, 'Flask – web framework and REST API')
    add_bullet(doc, 'Flask-CORS – cross-origin support for frontend')
    add_bullet(doc, 'Flask-SocketIO – WebSocket support for real-time events (e.g. notification_sent, high_risk_alert)')
    add_bullet(doc, 'Pandas – data loading, processing, and Excel I/O (openpyxl)')
    add_bullet(doc, 'python-dotenv – environment and .env configuration')
    add_bullet(doc, 'OpenAI (openai package) – GPT models for risk reasoning, sentiment, care routing, chatbot, and message generation')
    add_bullet(doc, 'Plotly (plotly) – server-side chart data generation (trend charts)')
    add_bullet(doc, 'Twilio – SMS, WhatsApp, and voice calls')
    add_bullet(doc, 'SendGrid (requests) – email delivery')
    add_bullet(doc, 'Pushover (requests) – push notifications to mobile app')
    add_bullet(doc, 'PyInstaller – optional packaging as standalone executable')

    add_heading(doc, '2.2 Frontend', level=2)
    add_bullet(doc, 'React 18 – UI framework')
    add_bullet(doc, 'React Plotly.js / react-plotly.js – interactive charts')
    add_bullet(doc, '@amcharts/amcharts5 – additional charting (e.g. gauges)')
    add_bullet(doc, 'Axios – HTTP client for API calls')
    add_bullet(doc, 'socket.io-client – WebSocket client for real-time updates')
    add_bullet(doc, 'react-app-rewired – build configuration (e.g. for Plotly)')
    add_bullet(doc, 'CSS – styling (e.g. Dashboard.css, KPICards.css, Charts.css, DataSnapshot.css, Chatbot.css, AgentsFlow.css)')

    add_heading(doc, '2.3 Data and configuration', level=2)
    add_bullet(doc, 'Excel (.xlsx) as primary data source – patient records, adherence %, refill days, contact info, consent, notification timestamps')
    add_bullet(doc, 'Config via .env: OPENAI_API_KEY, OPENAI_MODEL, EXCEL_PATH, Twilio, SendGrid, Pushover, thresholds (THRESHOLD_LOW, THRESHOLD_MED), OUTREACH_MAX_PER_DAY, OUTREACH_COOLDOWN_DAYS, etc.')
    add_bullet(doc, 'config.py – central config; runtime_paths.py – app/exe directory resolution for .env and Data folder')

    # 3. How agents are called
    add_heading(doc, '3. How the Agents Are Called', level=1)
    add_para(doc, 'All agent logic lives in agent.py. The backend (app_flask.py) and the background BackendRunner call these agents as follows.')

    add_heading(doc, '3.1 Background monitoring loop (BackendRunner._tick)', level=2)
    add_para(doc, 'A daemon thread runs periodically (AUTO_REFRESH_SECONDS, default 60). Each tick:')
    add_bullet(doc, 'Load patient data from Excel (utils.load_patient_data)')
    add_bullet(doc, 'For each row: validate_and_normalize_row (Data Quality Agent) → clean row')
    add_bullet(doc, 'For each clean row: assess_adherence_risk (Adherence Risk Agent)')
    add_bullet(doc, 'Build policy_context (today_sent, cap, cooldown_days, last_sent_ts)')
    add_bullet(doc, 'For each row: orchestrate_refill_and_notify(row, send=send_auto_notifications, policy_context=policy_context)')
    add_para(doc, 'Orchestration Agent internally calls: assess_adherence_risk, analyze_patient_sentiment_and_behavior, check_and_appreciate_refilled_patient, determine_care_routing, identify_patient_care_needs, build_adherence_message, determine_contact_route, enforce_policy_and_throttle; then, if send and allowed, it calls notify_patient or Twilio/SendGrid (voice, SMS, WhatsApp, email) or Pushover.')
    add_bullet(doc, 'On send: update_notification_sent_on writes back to Excel; socketio.emit("notification_sent", event) notifies connected clients')

    add_heading(doc, '3.2 API-driven agent calls', level=2)
    add_bullet(doc, 'GET /api/dashboard-data: loads data, then for each row runs validate_and_normalize_row and assess_adherence_risk; aggregates KPIs and high-risk list; returns JSON for frontend')
    add_bullet(doc, 'GET /api/patient-snapshot: same normalization and risk; optional filters; returns snapshot list')
    add_bullet(doc, 'POST /api/notification-preview: normalizes row, builds message (build_adherence_message), determine_contact_route; returns preview and channel')
    add_bullet(doc, 'POST /api/send-notification: resolve patient, normalize, orchestrate_refill_and_notify(..., send=True) for chosen channel (call/SMS/WhatsApp/email/Pushover)')
    add_bullet(doc, 'POST /api/make-call: make_twilio_voice_call with message from build_adherence_message')
    add_bullet(doc, 'POST /api/patient-response: analyze_patient_response (OpenAI), optional generate_resolution_recommendations, optional auto_respond_to_patient; returns analysis and suggested reply')
    add_bullet(doc, 'POST /api/chatbot: chatbot_agent (OpenAI) with conversation history; returns assistant reply')
    add_bullet(doc, 'POST /api/patient-query: patient lookup and risk/adherence info for dashboard/UI')

    add_heading(doc, '3.3 Agent list (agent.py)', level=2)
    add_bullet(doc, 'Data Quality Agent: validate_and_normalize_row – normalizes columns, consent, refill_due, city, NotificationSentOn, etc.')
    add_bullet(doc, 'Adherence Risk Agent: assess_adherence_risk – risk score/label from adherence %, thresholds, refill days')
    add_bullet(doc, 'Consent & Channel Routing Agent: determine_contact_route – allowed channel (SMS, email, call, WhatsApp, Pushover) and address from consent and NotificationChannel')
    add_bullet(doc, 'Safety & Throttling Agent: enforce_policy_and_throttle – daily cap, cooldown, one contact per patient per day')
    add_bullet(doc, 'Orchestration Agent: orchestrate_refill_and_notify – plans actions (notify_patient, appreciate_patient, escalate_care_team, send_nudge), applies route and policy, optionally sends notification')
    add_bullet(doc, 'Sentiment/Behavior: analyze_patient_sentiment_and_behavior – OpenAI-based analysis')
    add_bullet(doc, 'Appreciation: check_and_appreciate_refilled_patient – decides if to send appreciation (e.g. refill improved)')
    add_bullet(doc, 'Care Routing: determine_care_routing – whether to route to clinician; identify_patient_care_needs – care needs from risk/sentiment')
    add_bullet(doc, 'Patient response: analyze_patient_response, generate_resolution_recommendations, auto_respond_to_patient – OpenAI')
    add_bullet(doc, 'Chatbot: chatbot_agent – conversational AI (OpenAI) with system prompt and history')
    add_bullet(doc, 'Helpers: build_adherence_message (includes pharmacy table from pharmacy.py), notify_patient, send_pushover_message, send_twilio_sms, send_twilio_whatsapp, send_email, make_twilio_voice_call')

    # 4. End-to-end flow
    add_heading(doc, '4. End-to-End Flow', level=1)
    add_para(doc, 'Startup: main.py loads config (config.py loads .env from app dir / CWD), then starts Flask + SocketIO (app_flask). BackendRunner starts and runs _tick every AUTO_REFRESH_SECONDS.')
    add_para(doc, 'Data flow: Excel (EXCEL_PATH) → load_patient_data → per-row validate_and_normalize_row → assess_adherence_risk → orchestrate_refill_and_notify (with determine_contact_route, enforce_policy_and_throttle). If send and allowed, notification is sent via Twilio/SendGrid/Pushover and NotificationSentOn (and optional channel/message id) written back to Excel.')
    add_para(doc, 'Dashboard: Frontend (React) calls GET /api/dashboard-data and GET /api/trend-chart; subscribes to SocketIO for notification_sent / high_risk_alert. Renders KPIs, charts, high-risk table, and optional Data Snapshot, Chatbot, System Status.')
    add_para(doc, 'Manual actions: User can trigger notification preview, send notification, make call, or use chatbot/patient-response APIs from UI; those endpoints call the same agents (orchestrate_refill_and_notify, make_twilio_voice_call, chatbot_agent, analyze_patient_response, etc.).')

    # 5. Design and functionality
    add_heading(doc, '5. Design and Functionality', level=1)
    add_para(doc, 'Backend design:')
    add_bullet(doc, 'Single Flask app serves REST API, static frontend build, and SocketIO; CORS enabled for dev (frontend on different port)')
    add_bullet(doc, 'State: backend_runner (BackendRunner), live_on, send_auto_notifications, refresh_interval, audit_log, high_risk_notifications, patient_previous_states')
    add_bullet(doc, 'Care teams: persisted in Data/care_teams.json; API /api/care-teams for list/create and get by id')
    add_para(doc, 'Frontend design:')
    add_bullet(doc, 'Components: App.js, Dashboard, KPICards, Charts, ExecutiveSummary, ImportantDetails, DataSnapshot, Chatbot, ToastNotification, SystemStatus, AgentsFlow, Header')
    add_bullet(doc, 'Dashboard fetches /api/dashboard-data and /api/trend-chart; DataSnapshot uses /api/patient-snapshot, /api/notification-preview, /api/send-notification; Chatbot uses /api/chatbot')
    add_bullet(doc, 'Real-time: socket.io-client connects to backend; listens for notification_sent and high_risk_alert; can show toasts or refresh data')
    add_bullet(doc, 'Styling: component-level CSS files; responsive layout and animations for KPIs/charts')
    add_para(doc, 'Data design:')
    add_bullet(doc, 'utils.ensure_excel_columns ensures required columns exist (Member ID, Patient Name, Adherence %, Refill Days, contact, consent, NotificationSentOn, etc.); utils.update_notification_sent_on updates Excel after sends')
    add_bullet(doc, 'pharmacy.py: city → nearest pharmacy (static map + optional OpenStreetMap); messages include pharmacy table for refill guidance')

    # 6. Key files
    add_heading(doc, '6. Key Project Files', level=1)
    add_bullet(doc, 'main.py – entry point; startup checks; port detection; runs Flask + SocketIO')
    add_bullet(doc, 'app_flask.py – Flask app, CORS, SocketIO; routes (health, dashboard-data, trend-chart, system-status, system-control, patient-snapshot, notification-preview, send-notification, make-call, patient-response, chatbot, care-teams); BackendRunner; serve SPA')
    add_bullet(doc, 'agent.py – all agents and OpenAI/Twilio/SendGrid/Pushover helpers')
    add_bullet(doc, 'config.py – load .env; expose OPENAI_*, TWILIO_*, SENDGRID_*, PUSHOVER_*, EXCEL_PATH, thresholds, etc.')
    add_bullet(doc, 'utils.py – load_patient_data, ensure_excel_columns, update_notification_sent_on')
    add_bullet(doc, 'pharmacy.py – pharmacy lookup by city; format_pharmacy_table for messages')
    add_bullet(doc, 'structure.py – HTML template helpers, formatting (e.g. trend, name_parts)')
    add_bullet(doc, 'runtime_paths.py – get_app_dir, get_resource_dir for exe vs dev')
    add_bullet(doc, 'frontend/src/App.js – root; Dashboard and routing')
    add_bullet(doc, 'frontend/src/components/*.js – Dashboard, KPICards, Charts, DataSnapshot, Chatbot, etc.')
    add_bullet(doc, 'requirements.txt – Flask, flask-cors, flask-socketio, pandas, openai, plotly, twilio, python-dotenv, openpyxl, etc.')
    add_bullet(doc, 'frontend/package.json – React, react-plotly.js, axios, socket.io-client, @amcharts/amcharts5')

    add_heading(doc, '7. Summary', level=1)
    add_para(doc, 'The project is a full-stack Medication Adherence Dashboard: Flask + SocketIO backend with an agentic AI pipeline (data quality, risk, consent routing, throttling, orchestration, sentiment, care routing, and chatbot). The frontend is React with Plotly and AmCharts for visualizations. Agents are invoked from the background monitoring loop and from REST endpoints for dashboard, snapshot, notifications, calls, and chatbot. Notifications are sent via Twilio (SMS, WhatsApp, voice), SendGrid (email), or Pushover; patient data is stored in Excel and optionally packaged as a standalone executable with PyInstaller.')

    out_path = "Project_Development_Details.docx"
    doc.save(out_path)
    print("Created:", out_path)

if __name__ == "__main__":
    main()
